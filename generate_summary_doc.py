"""
Generate the portfolio project summary Word document.
Run with: python generate_summary_doc.py
Produces:  "Algo Trading Portfolio — Project Summary.docx"
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1A, 0x3A, 0x5C)
BLUE       = RGBColor(0x21, 0x96, 0xF3)
GREEN      = RGBColor(0x2E, 0x7D, 0x32)
ORANGE     = RGBColor(0xE6, 0x51, 0x00)
LIGHT_BLUE = RGBColor(0xE3, 0xF2, 0xFD)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY   = RGBColor(0xE0, 0xE0, 0xE0)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x21, 0x21, 0x21)

GITHUB_URL = "https://github.com/joe-bennett/first-algo-and-backtest"
ALPACA_URL = "https://alpaca.markets"
SIMFIN_URL = "https://simfin.com"


# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)  # RGBColor.__str__ returns 'RRGGBB'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)  # RGBColor.__str__ returns 'RRGGBB'
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2196F3')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(color)
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return hyperlink


def add_toc(doc):
    """Insert a Word TOC field that auto-populates when opened in Word."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'Right-click this text and choose "Update Field" to generate the Table of Contents.'
    placeholder.append(t)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)


def page_break(doc):
    doc.add_page_break()


# ── Style helpers ─────────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(20)
        run.font.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.size = Pt(14)
        run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(12)
        run.font.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold_phrases=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_phrases:
        remaining = text
        for phrase in bold_phrases:
            idx = remaining.find(phrase)
            if idx == -1:
                continue
            if idx > 0:
                run = p.add_run(remaining[:idx])
                run.font.size = Pt(11)
                run.font.color.rgb = BLACK
            bold_run = p.add_run(phrase)
            bold_run.bold = True
            bold_run.font.size = Pt(11)
            bold_run.font.color.rgb = BLACK
            remaining = remaining[idx + len(phrase):]
        if remaining:
            run = p.add_run(remaining)
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def callout(doc, title, text, bg=LIGHT_BLUE, title_color=NAVY):
    """A highlighted info box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6)
    if title:
        tp = cell.add_paragraph()
        tr = tp.add_run(f"  {title}")
        tr.bold = True
        tr.font.color.rgb = title_color
        tr.font.size = Pt(11)
        tp.paragraph_format.space_after = Pt(2)
    bp = cell.add_paragraph()
    br = bp.add_run(f"  {text}")
    br.font.size = Pt(11)
    br.font.color.rgb = BLACK
    bp.paragraph_format.space_after = Pt(4)
    # Remove the default empty first paragraph in cell
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    doc.add_paragraph()  # spacer


def code_block(doc, text):
    p = doc.add_paragraph()
    set_para_bg(p, LIGHT_GREY)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows, header_bg=NAVY, header_fg=WHITE, alt_bg=LIGHT_GREY):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = header_fg
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = alt_bg if r_idx % 2 == 0 else WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# Default paragraph font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Algorithmic Trading Portfolio")
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = NAVY

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run("A Quantitative Equity Strategy — Project Summary")
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = BLUE
sub_run.italic = True

doc.add_paragraph()
doc.add_paragraph()

divider_p = doc.add_paragraph()
divider_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
divider_run = divider_p.add_run("─" * 50)
divider_run.font.color.rgb = MID_GREY

doc.add_paragraph()

desc_p = doc.add_paragraph()
desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
desc_run = desc_p.add_run(
    "A fully automated stock-picking and portfolio management system built from scratch.\n"
    "Driven by academic factor research. Backed by rigorous historical testing.\n"
    "Live paper trading on Alpaca with email alerts and an interactive dashboard."
)
desc_run.font.size = Pt(12)
desc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f"Prepared: {datetime.date.today().strftime('%B %Y')}")
date_run.font.size = Pt(11)
date_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

github_p = doc.add_paragraph()
github_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
github_p.add_run("GitHub Repository:  ").font.size = Pt(11)
add_hyperlink(github_p, GITHUB_URL, GITHUB_URL)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tl_run = tagline_p.add_run('"Letting math and data do the stock-picking."')
tl_run.italic = True
tl_run.font.size = Pt(13)
tl_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

toc_heading = doc.add_heading("Table of Contents", level=1)
for run in toc_heading.runs:
    run.font.color.rgb = NAVY
    run.font.size = Pt(20)

callout(doc, "How to use this Table of Contents",
        "When you open this document in Microsoft Word, right-click anywhere in the "
        "grey table of contents area below and select 'Update Field' → 'Update entire table'. "
        "Word will generate a fully clickable, hyperlinked table of contents automatically.",
        bg=LIGHT_BLUE, title_color=NAVY)

add_toc(doc)
page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT IS THIS PROJECT?
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "1. What Is This Project?")

body(doc,
     "This project is a fully automated algorithmic trading system — a piece of software "
     "that uses mathematics, data, and academic research to decide which stocks to buy, "
     "which to sell short, and when to do each. It removes emotion and guesswork from "
     "investing and replaces them with a rules-based, repeatable process.",
     bold_phrases=["fully automated algorithmic trading system", "removes emotion and guesswork"])

body(doc,
     "Think of it this way: instead of watching CNBC, reading analyst reports, and making "
     "gut-feel decisions, this system reads financial data on every company in the S&P 500, "
     "scores each one on objective criteria, ranks them from best to worst, and builds a "
     "portfolio from the top-ranked and bottom-ranked stocks. Every quarter it rebalances "
     "automatically. Every evening it scans for new opportunities and sends an email alert.")

h2(doc, "What Problem Does It Solve?")

body(doc,
     "The vast majority of actively managed mutual funds underperform a simple S&P 500 "
     "index fund over the long run. The reason is not that fund managers are unintelligent "
     "— it is that human judgment in investing is reliably distorted by emotions: fear "
     "during crashes, greed during bubbles, overconfidence after wins, and loss aversion "
     "after losses. Academic research going back to the 1960s has consistently found that "
     "a small set of measurable characteristics predict which stocks will outperform. "
     "This system is built to exploit those characteristics systematically.",
     bold_phrases=["vast majority of actively managed mutual funds underperform",
                   "systematically"])

h2(doc, "What Did We Build?")

bullet(doc, "A stock ranking engine that scores every S&P 500 company on value, momentum, and quality factors")
bullet(doc, "A 120/20 long-short equity strategy: long the best-ranked stocks, short the worst")
bullet(doc, "An iron condor options scanner for opportunistic premium collection")
bullet(doc, "A backtesting framework that runs 10+ years of simulated history with institutional-quality bias corrections")
bullet(doc, "A live Streamlit dashboard with portfolio overview, backtest runner, signal scanner, and research sandbox")
bullet(doc, "Automated paper trading on Alpaca: rebalances quarterly, manages stop-losses daily")
bullet(doc, "An email alert system that explains every trade recommendation in plain English")
bullet(doc, "Advanced modes: concentrated positioning, conviction-weighted sizing, put options on the short book")

body(doc,
     "Everything is configurable through a single YAML file — no coding required to change "
     "strategy parameters. The entire codebase is open-source and available on GitHub:")

link_p = doc.add_paragraph()
add_hyperlink(link_p, GITHUB_URL, GITHUB_URL)
link_p.paragraph_format.space_after = Pt(6)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — INVESTMENT PHILOSOPHY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "2. The Investment Philosophy")

h2(doc, "Why Quantitative (Factor) Investing?")

body(doc,
     "Academic researchers — most notably Eugene Fama and Kenneth French, who won the Nobel "
     "Prize in Economics for this work — discovered that stocks with certain measurable "
     "characteristics consistently outperform the market over long periods. These "
     "characteristics are called factors. The most robustly documented are:",
     bold_phrases=["Eugene Fama and Kenneth French", "Nobel Prize", "factors"])

bullet(doc, "Value — cheap stocks (low price relative to earnings, assets, or cash flow) outperform expensive ones")
bullet(doc, "Momentum — stocks that have risen over the past year tend to keep rising; fallers tend to keep falling")
bullet(doc, "Quality — profitable, low-debt, efficiently-run companies outperform financially weak ones")

body(doc,
     "These findings have been replicated across dozens of countries, time periods, and "
     "asset classes. They work not because the market is irrational in an obvious way, but "
     "because they require patience and discipline that most investors cannot maintain — "
     "cheap stocks often feel uncomfortable to buy precisely because they have bad news "
     "priced in. A rules-based system has no such discomfort.")

h2(doc, "Why Long-Short (120/20)?")

body(doc,
     "A traditional portfolio only makes money when its stocks go up. A long-short portfolio "
     "makes money in two ways: when the longs go up AND when the shorts go down. This is "
     "called capturing alpha on both sides of the factor.",
     bold_phrases=["long-short portfolio", "capturing alpha on both sides"])

body(doc,
     "The '120/20' structure specifically means:")
bullet(doc, "120% of the portfolio is invested long (in the best-ranked stocks)")
bullet(doc, "20% of the portfolio is sold short (in the worst-ranked stocks)")
bullet(doc, "The short sale proceeds fund the extra 20% of long buying — so the net exposure is still 100%")

callout(doc, "Concrete Example",
        "With a $100,000 portfolio:\n"
        "  • Short $20,000 of the worst-ranked stocks → receive $20,000 in cash proceeds\n"
        "  • Use that $20,000 + original $100,000 to buy $120,000 of the best-ranked stocks\n"
        "  • Net investment: still $100,000. Gross exposure: $140,000.\n"
        "  • If longs rise 10% and shorts fall 5%: profit = $12,000 + $1,000 = $13,000 (13% return)",
        bg=LIGHT_BLUE)

h2(doc, "Why the S&P 500?")

body(doc,
     "The S&P 500 is a list of the 500 largest publicly traded US companies — Apple, "
     "Microsoft, ExxonMobil, JPMorgan, and 496 others. We use it as our universe because:")
bullet(doc, "All companies are highly liquid — easy to buy and sell in any size")
bullet(doc, "Short selling is practical — shares are easy to borrow")
bullet(doc, "Historical membership data is available going back to 1996 — critical for unbiased backtesting")
bullet(doc, "Data quality is high — fundamentals data is reliable for all 500 companies")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE 120/20 STRATEGY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "3. Strategy 1 — The 120/20 Value-Momentum-Quality System")

h2(doc, "How Stocks Get Scored")

body(doc,
     "Every stock in the S&P 500 receives a composite score between 0 and 1. A score near "
     "1.0 means the stock ranks among the best in the universe on the combined factors. "
     "A score near 0.0 means it ranks among the worst. The top 20% of scorers become the "
     "long book; the bottom 20% become the short book.",
     bold_phrases=["composite score between 0 and 1"])

body(doc,
     "The score is built from three factor categories, each explained below.")

h2(doc, "Factor 1: Value (40% of final score)")

body(doc,
     "Value investing — popularised by Benjamin Graham and Warren Buffett — is the idea "
     "that buying a dollar of earnings or assets for fifty cents is a good deal. "
     "We measure value four ways:",
     bold_phrases=["Benjamin Graham and Warren Buffett"])

add_table(doc,
    headers=["Factor", "What It Measures", "Direction"],
    rows=[
        ["P/E Ratio", "Price paid for every $1 of annual earnings", "Lower = better value"],
        ["P/B Ratio", "Price paid for every $1 of net assets (book value)", "Lower = better value"],
        ["FCF Yield", "Free cash flow generated per $1 of market cap", "Higher = more cash for every dollar paid"],
        ["EV/EBITDA", "Total cost of the business vs. operating earnings", "Lower = better value"],
    ])

callout(doc, "Plain English Example",
        "Imagine two companies both earn $10 per share annually.\n"
        "  Company A trades at $100/share → P/E of 10 (you pay $10 for every $1 of earnings)\n"
        "  Company B trades at $250/share → P/E of 25 (you pay $25 for every $1 of earnings)\n"
        "Company A has a much higher value score — you're getting the same $1 of earnings "
        "for a much lower price.", bg=LIGHT_BLUE)

h2(doc, "Factor 2: Momentum (40% of final score)")

body(doc,
     "Momentum is one of the most robustly documented phenomena in finance: stocks that "
     "have performed well over the past year tend to continue outperforming, and stocks "
     "that have performed poorly tend to continue underperforming. This is sometimes called "
     "the 'trend following' effect.",
     bold_phrases=["Momentum", "most robustly documented phenomena in finance"])

body(doc,
     "We measure momentum as the 12-month price return, skipping the most recent month. "
     "Skipping the last month avoids a well-known short-term reversal effect where recent "
     "winners temporarily give back gains. The calculation uses 5-day price averages at "
     "each reference point — rather than a single day's price — to prevent one unusually "
     "volatile day from distorting an entire year of price history.",
     bold_phrases=["12-month price return, skipping the most recent month",
                   "5-day price averages"])

h2(doc, "Factor 3: Quality (20% of final score)")

body(doc,
     "Quality factors distinguish financially strong, well-run businesses from weak, "
     "over-leveraged ones. We measure quality three ways:",
     bold_phrases=["Quality factors"])

add_table(doc,
    headers=["Factor", "What It Measures", "Direction"],
    rows=[
        ["Return on Equity (ROE)", "Net income earned per $1 of shareholder capital", "Higher = management generates more profit"],
        ["Net Profit Margin", "Fraction of each dollar of revenue that becomes profit", "Higher = more efficient business"],
        ["Debt/Equity Ratio", "Total debt relative to shareholder equity", "Lower = less financial risk"],
    ])

h2(doc, "The Percentile Ranking System")

body(doc,
     "A critical technical detail: raw numbers cannot be directly compared across factors "
     "or across time. A P/E of 15 might be cheap in 2022 but expensive in 2010. A tech "
     "company at P/E 25 might be cheap for tech but expensive for a utility. "
     "The solution is percentile ranking.",
     bold_phrases=["percentile ranking"])

body(doc,
     "Instead of using the raw number, each stock is ranked relative to every other stock "
     "in the universe today. A stock's P/E percentile score answers the question: "
     "'What fraction of S&P 500 companies are more expensive than this one right now?'")

add_table(doc,
    headers=["Stock", "Raw P/E", "Rank Among 500", "Percentile", "Value Sub-Score (inverted)"],
    rows=[
        ["ExxonMobil (XOM)", "11", "8th cheapest", "1.6%", "0.98 ← near perfect"],
        ["Apple (AAPL)", "28", "350th cheapest", "70%", "0.30"],
        ["Tesla (TSLA)", "55", "470th cheapest", "94%", "0.06 ← very expensive"],
    ])

body(doc,
     "Because every factor is converted to a 0–1 scale before combining, the final "
     "composite score is always comparable regardless of market conditions. A 0.85 score "
     "always means 'in the top 15% of the universe on the combined signal today.'",
     bold_phrases=["0.85 score always means 'in the top 15% of the universe'"])

h2(doc, "The Final Composite Score")

body(doc, "The three category scores combine into one final ranking:")
code_block(doc,
    "Final Score  =  0.40 × Value Score\n"
    "             +  0.40 × Momentum Score\n"
    "             +  0.20 × Quality Score\n\n"
    "Top 20% of final scores  →  Long Book (buy these)\n"
    "Bottom 20% of scores     →  Short Book (sell these short)")

body(doc,
     "All weights are configurable in a single YAML file. Running a backtest after any "
     "change shows the historical impact immediately — no guesswork required.")

h2(doc, "Sector Neutralization")

body(doc,
     "By default, stocks compete globally for rankings. Optionally, sector neutralization "
     "forces stocks to compete only within their own industry sector. This prevents the "
     "entire long book from filling up with one cheap sector (e.g., all energy companies "
     "when energy is out of favor). The default is global ranking because sector cheapness "
     "vs. expensiveness is itself a valid signal.",
     bold_phrases=["sector neutralization"])

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — IRON CONDOR STRATEGY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "4. Strategy 2 — The Iron Condor Options Scanner")

h2(doc, "What Are Options?")

body(doc,
     "An option is a contract that gives the buyer the right — but not the obligation — "
     "to buy or sell 100 shares of a stock at a specific price before a specific date. "
     "Options have a price (called the premium). The seller of the option collects that "
     "premium upfront and keeps it if the option expires without being used.",
     bold_phrases=["option", "premium"])

h2(doc, "What Is an Iron Condor?")

body(doc,
     "An iron condor is a four-legged options trade that profits when a stock stays within "
     "a range. You simultaneously sell a call spread above the current price and a put "
     "spread below it. As long as the stock stays between those two boundaries by "
     "expiration, you keep all the premium collected.",
     bold_phrases=["iron condor", "profits when a stock stays within a range"])

callout(doc, "Iron Condor Example",
        "Stock XYZ is trading at $100. You set up a condor:\n"
        "  Downside protection: Buy $85 put / Sell $90 put\n"
        "  Upside protection:   Sell $110 call / Buy $115 call\n\n"
        "  You collect $2.50/share in premium = $250 per contract.\n"
        "  As long as XYZ stays between $90 and $110 by expiration, you keep the $250.\n"
        "  Max loss: $250 (if stock breaks out strongly in either direction).",
        bg=LIGHT_BLUE)

h2(doc, "When Does the System Recommend a Condor?")

body(doc,
     "The scanner fires when two conditions are met:")
bullet(doc, "IV Rank ≥ 50 — Implied Volatility is in the top half of its range over the past year. This means options are expensive relative to their recent history, making it a good time to sell premium.")
bullet(doc, "The stock is NOT already held in the 120/20 equity book — a condor bets on range-bound movement, which directly conflicts with holding a directional long or short position in the same stock.")

body(doc,
     "The system automatically calculates how many contracts to trade based on portfolio "
     "size and sends a detailed email with the exact trade structure, credit received, "
     "maximum loss, and management rules.")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — DATA PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "5. The Data Pipeline")

body(doc,
     "A strategy is only as good as the data it runs on. We built a multi-source data "
     "pipeline that provides clean, point-in-time accurate data for both live signals "
     "and historical backtesting.")

h2(doc, "Universe: S&P 500 Membership")

body(doc,
     "The strategy runs on the S&P 500. The list is fetched automatically from Wikipedia. "
     "For backtesting, we use a historical membership dataset that tracks exactly which "
     "companies were in the S&P 500 on every day going back to 1996 — so the backtest "
     "only uses stocks that were actually in the index on each date, not today's survivors.")

h2(doc, "Price Data: yfinance")

body(doc,
     "Daily adjusted closing prices are downloaded from Yahoo Finance via the yfinance "
     "Python library. Prices are cached locally and refreshed every 4 hours during market "
     "hours to avoid redundant downloads. Price data is used for momentum calculations "
     "and for computing portfolio values.")

h2(doc, "Fundamental Data: SimFin (Backtesting)")

body(doc,
     "SimFin provides quarterly financial statement data for 3,600+ US companies going "
     "back to 2007. Critically, SimFin records both the fiscal period end date and the "
     "date the filing was made public — this allows the backtest to only use data that "
     "was actually available on each historical date, eliminating look-ahead bias.",
     bold_phrases=["look-ahead bias"])

body(doc, "Fundamental data from SimFin covers:")
bullet(doc, "Income statement: revenue, net income, operating earnings (EBITDA)")
bullet(doc, "Balance sheet: total assets, total debt, book value (equity)")
bullet(doc, "Cash flow statement: free cash flow")
bullet(doc, "Computed ratios: P/E, P/B, FCF yield, EV/EBITDA, ROE, net margin, debt/equity")

body(doc,
     "The SimFin dataset (~450MB) is downloaded once and cached locally. It auto-refreshes "
     "every 7 days. A free SimFin API key is required (no credit card).")

h2(doc, "Fundamental Data: yfinance (Live Signals)")

body(doc,
     "For live signal scanning, fundamentals are pulled directly from Yahoo Finance via "
     "yfinance. This is faster than SimFin for real-time use. The data reflects the most "
     "recent available filings, which is appropriate for live trading decisions.")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — BACKTESTING
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "6. Backtesting — How We Validated the Strategy")

h2(doc, "What Is a Backtest?")

body(doc,
     "A backtest is a historical simulation: take the strategy, run it on past data as if "
     "you were making decisions in real time, and see what would have happened. It answers "
     "the question: 'If this system existed in 2014, what would the returns have looked "
     "like through 2024?'",
     bold_phrases=["backtest", "historical simulation"])

body(doc,
     "Backtesting is the primary way to validate whether a strategy has genuine edge "
     "before putting real money behind it. However, backtests are notorious for producing "
     "misleading results if not built carefully. We implemented two institutional-quality "
     "bias corrections that most retail backtests skip entirely.",
     bold_phrases=["two institutional-quality bias corrections"])

h2(doc, "Bias Fix 1: Eliminating Look-Ahead Bias")

body(doc,
     "Look-ahead bias occurs when a backtest uses information that was not available at "
     "the time the decision was being made. The most common example: using today's P/E "
     "ratio to make a 2018 investment decision. A company's 2018 Q3 earnings report might "
     "not have been published until November 2018 — using it for a September 2018 decision "
     "is cheating.",
     bold_phrases=["Look-ahead bias", "using information that was not available"])

body(doc,
     "Our fix: SimFin records the exact date each quarterly filing was made public. The "
     "backtest only uses a filing after that public date. A March 31 quarter-end that was "
     "filed on May 15 is not available to the model until May 15 — not a day earlier. "
     "This adds complexity but produces realistic results.",
     bold_phrases=["only uses a filing after that public date"])

h2(doc, "Bias Fix 2: Eliminating Survivorship Bias")

body(doc,
     "Survivorship bias occurs when a backtest only includes companies that still exist "
     "today. If you run a 2010–2024 backtest using the current S&P 500 list, you are "
     "automatically including only companies that were successful enough to still be in "
     "the index in 2024. Companies that went bankrupt, were acquired at a loss, or were "
     "removed for poor performance are invisible — making historical returns look better "
     "than they actually were.",
     bold_phrases=["Survivorship bias", "companies that still exist today"])

callout(doc, "Analogy",
        "Imagine ranking restaurants based only on the ones still open in 2024. "
        "You would never see the ones that closed in 2016 due to poor food quality. "
        "Your ranking system would look like it only picked winners — because the losers "
        "are hidden from history.", bg=LIGHT_BLUE)

body(doc,
     "Our fix: we use a historical S&P 500 membership dataset (sourced from the fja05680 "
     "GitHub project) that records which companies were in the index on every single day "
     "going back to 1996. At each quarterly rebalance date in the backtest, only stocks "
     "that were actually in the S&P 500 on that date are eligible for ranking. Companies "
     "that were later removed, went bankrupt, or were acquired are correctly excluded from "
     "future dates but included in the periods when they were actual members.",
     bold_phrases=["historical S&P 500 membership dataset"])

h2(doc, "The Backtesting Engine: VectorBT")

body(doc,
     "We use VectorBT — a professional-grade Python backtesting library — to simulate "
     "the portfolio. It handles fractional share positions, transaction costs (10 basis "
     "points per trade), slippage (5 basis points), and correctly models a single shared "
     "portfolio rather than running N independent simulations.",
     bold_phrases=["VectorBT"])

body(doc, "Each backtest run produces:")
bullet(doc, "An equity curve: portfolio value over time vs. the S&P 500 benchmark (SPY)")
bullet(doc, "A drawdown chart: how far below the peak the portfolio fell at each point")
bullet(doc, "A rolling Sharpe ratio: risk-adjusted return over trailing 63-day windows")
bullet(doc, "A metrics table: annualized return, max drawdown, Sharpe ratio, Sortino ratio, win rate")
bullet(doc, "An HTML report saved locally that can be opened in any browser")

h2(doc, "Running a Backtest")

body(doc, "From the dashboard — no coding required:")
bullet(doc, "Open the Backtest page in the dashboard")
bullet(doc, "Set the date range, factor weights, exposure levels, concentration settings")
bullet(doc, "Click Run Backtest — results appear in 2–5 minutes")
bullet(doc, "Run again with different settings — both runs appear side by side for comparison")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — THE DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "7. The Interactive Dashboard")

body(doc,
     "The system includes a Streamlit web dashboard that runs locally in your browser. "
     "It has four pages, each serving a distinct purpose. Launch it by running "
     "'streamlit run dashboard/app.py' in the project folder, or by typing "
     "'start-dashboard' in PowerShell.",
     bold_phrases=["four pages"])

h2(doc, "Page 1: Portfolio Overview")

body(doc,
     "Shows the live state of the paper trading account, pulled directly from Alpaca in "
     "real time. Includes:")
bullet(doc, "Account summary: portfolio value, cash, buying power, total unrealized P&L")
bullet(doc, "Position count breakdown: number of longs vs. shorts")
bullet(doc, "Holdings weight bar chart: each position as a percentage of the portfolio")
bullet(doc, "P&L bar chart: unrealized gain/loss by position, color-coded green/red")
bullet(doc, "Detailed position tables: entry price, current price, stop-loss price, P&L")
bullet(doc, "Refresh button: pulls the latest data from Alpaca on demand")

h2(doc, "Page 2: Backtest Runner")

body(doc,
     "A fully interactive backtest environment where every strategy parameter can be "
     "adjusted with sliders and toggles — no coding required. Controls include:")
bullet(doc, "Date range and initial capital")
bullet(doc, "Value / momentum / quality factor weight sliders")
bullet(doc, "Long/short exposure sliders — move the '120' and '20' in 120/20 (e.g., test 130/30)")
bullet(doc, "Book size: what percentage of the universe to include in the long and short books")
bullet(doc, "Concentration mode: limit the book to a fixed stock count (e.g., top 15 longs)")
bullet(doc, "Conviction weighting: give more capital to higher-ranked stocks within the book")
bullet(doc, "Sector neutralization toggle")
bullet(doc, "Auto-labeled run names so the comparison table is self-documenting")

h2(doc, "Page 3: Signals")

body(doc,
     "Shows what the strategy recommends right now — today's actual rankings and trade "
     "suggestions based on current market data. Includes a 'Dry Run' mode that generates "
     "the full email text and signal analysis without sending anything. "
     "Also runs the iron condor scanner on demand.")

h2(doc, "Page 4: Research Sandbox")

body(doc,
     "A fast experimentation environment that generates signals in seconds (no full "
     "backtest needed). Adjust factor weights, concentration, and conviction settings "
     "and immediately see how the top-ranked stock list changes. Ideal for quick "
     "'what if I weighted momentum more heavily?' questions.")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — AUTOMATED TRADING
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "8. Automated Paper Trading with Alpaca")

h2(doc, "What Is Alpaca?")

body(doc,
     "Alpaca (alpaca.markets) is a commission-free brokerage that offers a free, open "
     "API — meaning software can connect to it and place trades automatically. It also "
     "provides paper trading: a fully simulated account loaded with $100,000 of fake "
     "money where automated strategies can be tested without any financial risk.",
     bold_phrases=["Alpaca", "commission-free", "paper trading"])

body(doc,
     "The system currently runs on Alpaca paper trading. Every feature works exactly as "
     "it would with real money — orders are placed, positions are tracked, stop-losses "
     "are active — but no actual capital is at risk. This phase is used to validate the "
     "system's real-world performance before going live.")

h2(doc, "Automated Quarterly Rebalancing")

body(doc,
     "On the first trading day of each quarter (January, April, July, October), the "
     "system automatically:")
bullet(doc, "Runs the full signal scan on the S&P 500")
bullet(doc, "Identifies which stocks should be in the long and short books")
bullet(doc, "Compares the target portfolio to the current holdings")
bullet(doc, "Closes positions that are no longer in the signal list")
bullet(doc, "Opens or adjusts positions to match the new targets")
bullet(doc, "Attaches a 15% stop-loss order to every new long position")
bullet(doc, "Sends a summary email listing every order placed")

body(doc,
     "This runs automatically via Windows Task Scheduler — no human action required.")

h2(doc, "Daily Stop-Loss Replacement")

body(doc,
     "Every evening, a separate process runs that:")
bullet(doc, "Compares expected holdings (from last rebalance) to actual Alpaca positions")
bullet(doc, "Detects any long positions that were stopped out during the day")
bullet(doc, "Re-ranks the S&P 500 to find the next-best candidate not already held")
bullet(doc, "Buys the replacement at the same portfolio weight as the original")
bullet(doc, "Attaches a new 15% stop-loss to the replacement")
bullet(doc, "Sends an email notification describing the replacement")

body(doc,
     "This ensures that when a stop-loss fires, the portfolio weight is immediately "
     "redeployed into the next-best opportunity rather than sitting idle in cash.",
     bold_phrases=["immediately redeployed into the next-best opportunity"])

h2(doc, "Stop-Loss Logic")

body(doc,
     "Every long position has a server-side stop order sitting at Alpaca at 15% below "
     "the entry price. This is a 'good till cancelled' order that executes automatically "
     "if the price drops to that level — even if the computer running this software is "
     "off. The stop is attached the moment the position is opened.",
     bold_phrases=["server-side stop order", "good till cancelled"])

callout(doc, "Example",
        "Buy AAPL at $200/share → stop-loss order placed at $170/share (15% below).\n"
        "If AAPL drops to $170 at any point, Alpaca automatically sells the position.\n"
        "The next evening, the system detects the stopped-out position and buys a replacement.",
        bg=LIGHT_BLUE)

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — EMAIL ALERT SYSTEM
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "9. The Email Alert System")

body(doc,
     "The system sends plain-English email alerts through Gmail for every significant "
     "event. No trading jargon — every alert explains what to do and why.")

h2(doc, "What Triggers an Alert?")

add_table(doc,
    headers=["Event", "When", "What the Email Contains"],
    rows=[
        ["Quarterly Rebalance", "1st of Jan / Apr / Jul / Oct", "Full long and short signal list with factor scores and trade instructions"],
        ["Iron Condor Opportunity", "Daily scan after market close", "Trade structure, credit received, max loss, exact contract count, management rules"],
        ["Stop-Loss Replacement", "Evening after a stop fires", "Which stock was stopped out, which stock replaced it, and why"],
        ["Risk Breach", "Any time a limit is hit", "Which limit was breached and recommended action"],
    ])

h2(doc, "What an Equity Signal Email Looks Like")

code_block(doc,
    "=== PORTFOLIO SIGNAL: 2026-03-01 ===\n\n"
    "Long book: 15 positions | Short book: 10 put contracts + 5 short shares\n\n"
    "--- TOP 5 LONGS ---\n"
    "LONG XYZ @ 8.0% of portfolio\n"
    "  Composite score: 0.91  (value: 0.88, momentum: 0.93, quality: 0.82)\n"
    "  P/E 12.4 | P/B 1.8 | FCF yield 6.2% | ROE 24.1% | Margin 18.3%\n"
    "  WHY: XYZ ranks highly because it is cheap on fundamentals and has\n"
    "       strong 12-month price momentum.\n"
    "  HOW: Buy market order at open\n")

body(doc,
     "Every alert includes the WHY (which factors are driving the recommendation) and "
     "the HOW (exactly what order to place). Even someone who has never traded before "
     "can read the email and know what to do.")

h2(doc, "Technical Setup")

body(doc, "The email system uses Gmail's standard SMTP protocol — no third-party services required. Setup requires:")
bullet(doc, "A Gmail account")
bullet(doc, "A 16-character Gmail App Password (generated in Google Account settings)")
bullet(doc, "Two lines in the .env credentials file")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — ADVANCED FEATURES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "10. Advanced Features")

body(doc,
     "Beyond the core 120/20 strategy, several advanced modes were built to give more "
     "control over the risk/return profile. All are toggleable — off by default, "
     "configurable through a single YAML file or dashboard sliders.")

h2(doc, "Concentration Mode")

body(doc,
     "By default the portfolio spreads across ~100 long positions at 1.2% each. "
     "Concentration mode instead focuses capital on only the highest-ranked stocks.",
     bold_phrases=["Concentration mode"])

add_table(doc,
    headers=["Setting", "Long Count", "Position Size", "Character"],
    rows=[
        ["Default (top 20%)", "~100 stocks", "~1.2% each", "Diversified — smooth returns, lower single-stock risk"],
        ["Top 15 longs", "15 stocks", "~8% each", "Aggressive — higher return potential, higher drawdown"],
        ["Top 10 longs", "10 stocks", "~12% each", "Very aggressive — maximum factor concentration"],
    ])

body(doc,
     "The rationale: the ranking signal is strongest at the extremes. The #1 ranked stock "
     "is more compelling than the #80 ranked stock. Concentration forces the portfolio to "
     "actually reflect those differences rather than diluting them across 100 names.",
     bold_phrases=["the ranking signal is strongest at the extremes"])

h2(doc, "Conviction-Weighted Position Sizing")

body(doc,
     "Even within a concentrated book, equal weighting treats rank #1 and rank #15 "
     "identically. Conviction weighting sizes each position proportional to its composite "
     "score — so the highest-scoring stock gets more capital than the lowest-scoring stock "
     "still in the book.",
     bold_phrases=["Conviction weighting"])

code_block(doc,
    "Position weight = (stock's composite score) / (sum of all selected scores) × total long exposure\n\n"
    "Example — top 5 longs, 120% total long exposure:\n"
    "  Stock A  score 0.94  →  0.94/4.25 × 120%  =  26.5%  ← gets the most capital\n"
    "  Stock B  score 0.88  →  0.88/4.25 × 120%  =  24.8%\n"
    "  Stock C  score 0.83  →  0.83/4.25 × 120%  =  23.4%\n"
    "  Stock E  score 0.79  →  0.79/4.25 × 120%  =  22.3%  ← gets the least")

h2(doc, "Puts on the Short Book")

body(doc,
     "Instead of short-selling shares for the highest-conviction bearish signals, the "
     "system can buy put options. A put option gives the right to sell 100 shares at a "
     "fixed price — it profits when the stock falls, but with a convex (non-linear) payoff.",
     bold_phrases=["put options", "convex (non-linear) payoff"])

add_table(doc,
    headers=["", "Short Selling Shares", "Buying Put Options"],
    rows=[
        ["Max profit", "100% (stock goes to zero)", "300–500%+ on a large down move"],
        ["Max loss", "Unlimited (stock can rise forever)", "Premium paid (fixed, known upfront)"],
        ["Short squeeze risk", "Yes — can be forced to cover at a loss", "None — options cannot be squeezed"],
        ["Borrow fees", "Yes — must pay to borrow shares", "None"],
        ["Time pressure", "None — can hold indefinitely", "Must be right within the option's timeframe"],
    ])

body(doc,
     "The system uses this for the most extreme bottom-ranked stocks — the ones with the "
     "strongest bearish signal — where a large move is more likely. Less extreme shorts "
     "use regular short selling where the time-decay disadvantage of puts is not worth paying.",
     bold_phrases=["most extreme bottom-ranked stocks"])

h2(doc, "Adjustable Exposure (120/20 → 130/30 → 150/50)")

body(doc,
     "The long and short exposure levels (the '120' and '20' in 120/20) are fully "
     "adjustable. The Backtest page has dedicated sliders and shows the resulting strategy "
     "name in real time: '130/30 — net 100%, gross 160%'. Multiple exposure levels can be "
     "backtested side by side for comparison.",
     bold_phrases=["fully adjustable"])

add_table(doc,
    headers=["Structure", "Long Exposure", "Short Exposure", "Gross Exposure", "Character"],
    rows=[
        ["100/0", "100%", "0%", "100%", "Long-only — simplest, no shorting"],
        ["120/20", "120%", "20%", "140%", "Default — balanced L/S"],
        ["130/30", "130%", "30%", "160%", "More aggressive on both sides"],
        ["150/50", "150%", "50%", "200%", "High-octane — maximum factor exposure"],
    ])

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "11. The Technology Stack")

body(doc,
     "Every component of this system is built on free, open-source software. "
     "No paid data subscriptions are required beyond a free SimFin API key.")

add_table(doc,
    headers=["Technology", "Role", "Cost"],
    rows=[
        ["Python 3.10+", "Core programming language", "Free"],
        ["yfinance", "Live and historical price data; option chains", "Free"],
        ["SimFin", "Point-in-time historical fundamentals for backtesting", "Free (API key required)"],
        ["VectorBT", "Professional-grade portfolio backtesting engine", "Free"],
        ["Streamlit", "Interactive web dashboard", "Free"],
        ["Alpaca", "Paper and live trading API; order execution", "Free (paper trading)"],
        ["alpaca-py", "Python SDK for Alpaca API", "Free"],
        ["Plotly", "Interactive charts in the dashboard", "Free"],
        ["pandas / numpy", "Data manipulation and numerical computing", "Free"],
        ["PyYAML", "Reading configuration files", "Free"],
        ["python-dotenv", "Secure credential management via .env files", "Free"],
        ["smtplib (built-in)", "Email alerts via Gmail SMTP", "Free"],
        ["GitHub", "Version control and public code repository", "Free"],
    ])

body(doc,
     "The entire system runs on a standard Windows laptop. No cloud servers, no paid APIs, "
     "no ongoing subscription costs beyond whatever brokerage fees apply when trading live.",
     bold_phrases=["No cloud servers, no paid APIs, no ongoing subscription costs"])

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — HOW TO GET STARTED
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "12. How to Get Started")

body(doc,
     "Anyone can download and run this system. Here is the complete setup process from zero.")

h2(doc, "Prerequisites")
bullet(doc, "A Windows, Mac, or Linux computer")
bullet(doc, "Python 3.10 or later — download from python.org (free). Check 'Add Python to PATH' during installation on Windows.")
bullet(doc, "Git — download from git-scm.com (free). Used to download the code.")
bullet(doc, "A Gmail account (for email alerts)")
bullet(doc, "A free SimFin account at simfin.com (for backtesting fundamentals)")
bullet(doc, "Optional: A free Alpaca account at alpaca.markets (for paper trading automation)")

h2(doc, "Step-by-Step Setup")

body(doc, "Step 1 — Download the code:")
code_block(doc,
    "git clone https://github.com/joe-bennett/first-algo-and-backtest.git\n"
    "cd first-algo-and-backtest")

body(doc, "Step 2 — Install required Python packages:")
code_block(doc, "pip install -r requirements.txt")

body(doc, "Step 3 — Set up credentials. Copy .env.example to .env and fill in:")
code_block(doc,
    "GMAIL_ADDRESS=you@gmail.com\n"
    "GMAIL_APP_PASSWORD=xxxx xxxx xxxx xxxx\n"
    "ALERT_TO_EMAIL=you@gmail.com\n"
    "SIMFIN_API_KEY=your_simfin_key\n"
    "ALPACA_API_KEY=your_key          ← optional, leave blank if skipping Alpaca\n"
    "ALPACA_SECRET_KEY=your_secret    ← optional\n"
    "ALPACA_BASE_URL=https://paper-api.alpaca.markets  ← optional")

body(doc, "Step 4 — Verify the install:")
code_block(doc, "python quickstart.py")

body(doc, "Step 5 — Launch the dashboard:")
code_block(doc, "streamlit run dashboard/app.py")

body(doc,
     "The dashboard opens automatically in your browser. Go to the Backtest page first — "
     "run a historical simulation with default settings to see how the strategy would have "
     "performed. Explore the Research Sandbox to understand how factor weights affect "
     "stock selection. Run a signal scan to see today's recommendations.")

h2(doc, "What You Can Do Without an Alpaca Account")

body(doc,
     "Alpaca is entirely optional. The following work with zero brokerage account:")
bullet(doc, "All backtesting — full historical simulation with charts and metrics")
bullet(doc, "Research Sandbox — explore factor weights and ranking")
bullet(doc, "Signal scanning — see today's recommended trades")
bullet(doc, "Email alerts — get notified when signals or condor opportunities arise")

body(doc,
     "The only things that require Alpaca are the Portfolio Overview page (live positions), "
     "automated rebalancing, and automated stop-loss management. These are Phase 2 features "
     "— the analytical value of the system is fully accessible without them.")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — WHAT'S NEXT
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "13. What's Next — The Road to Live Trading")

add_table(doc,
    headers=["Phase", "Status", "Description"],
    rows=[
        ["Phase 1", "Complete", "Data pipeline, signal generation, backtesting, email alerts. Full analytical system built."],
        ["Phase 2", "Current", "Alpaca paper trading connected. Automated quarterly rebalancing, daily stop-loss replacement, portfolio ledger, live dashboard."],
        ["Phase 3", "Planned", "Live trading with real money. Additional risk controls, circuit breakers, position monitoring, and live P&L tracking against benchmark."],
    ])

h2(doc, "Before Going Live")

body(doc, "The paper trading phase (Phase 2) serves several important purposes:")
bullet(doc, "Validates that orders execute as expected in real market conditions")
bullet(doc, "Surfaces any bugs in the rebalancing or stop-loss logic before real money is involved")
bullet(doc, "Builds a live track record that can be compared against the backtest")
bullet(doc, "Gives the operator time to get comfortable watching the system make decisions autonomously")

body(doc,
     "When Phase 3 begins, the only change required is updating the Alpaca URL from "
     "the paper trading endpoint to the live endpoint and funding a live account. "
     "All code, strategies, and automation remain identical.",
     bold_phrases=["the only change required"])

h2(doc, "Potential Future Enhancements")
bullet(doc, "Kelly criterion position sizing — mathematically optimal bet sizing based on historical win rates")
bullet(doc, "Machine learning overlay — use gradient boosting to predict which factor signals are most reliable in the current market regime")
bullet(doc, "Live options management — automated closing of iron condor and put positions at profit targets")
bullet(doc, "Multi-strategy portfolio — run the value-momentum and condor strategies within a unified risk budget")
bullet(doc, "Real-time intraday monitoring — alerts if any position moves more than X% intraday")

page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, "14. Project Links and Resources")

add_table(doc,
    headers=["Resource", "Link", "Notes"],
    rows=[
        ["GitHub Repository", GITHUB_URL, "Full source code, README, GUIDE"],
        ["Alpaca Markets", ALPACA_URL, "Free paper trading account signup"],
        ["SimFin", SIMFIN_URL, "Free fundamentals API key"],
        ["Dashboard", "streamlit run dashboard/app.py", "Run locally after setup"],
    ])

doc.add_paragraph()
final_p = doc.add_paragraph()
final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = final_p.add_run(
    "This project represents a complete, institutional-quality algorithmic trading system "
    "built entirely from open-source tools and free data sources. From factor research "
    "to live paper trading, every component was designed, built, tested, and documented "
    "from scratch."
)
fr.italic = True
fr.font.size = Pt(11)
fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
final_p.paragraph_format.space_before = Pt(12)

doc.add_paragraph()
sign_p = doc.add_paragraph()
sign_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sign_p.add_run("Built with Python, patience, and a lot of factor research.")
sr.bold = True
sr.font.size = Pt(12)
sr.font.color.rgb = NAVY


# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "Algo Trading Portfolio — Project Summary.docx"
doc.save(output_path)
print(f"\nDocument saved: {output_path}")
print("Open in Microsoft Word, then right-click the Table of Contents and select")
print("'Update Field' -> 'Update entire table' to generate the hyperlinked TOC.")
